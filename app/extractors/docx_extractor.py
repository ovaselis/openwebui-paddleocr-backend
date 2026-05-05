import logging
import shutil
import zipfile
from pathlib import Path

from docx import Document
from fastapi import HTTPException, status

from app.config import settings
from app.extractors.image_extractor import validate_image
from app.ocr import get_engine
from app.schemas import OCRPage


logger = logging.getLogger("paddleocr_backend")


async def extract_docx(docx_path: Path) -> tuple[list[OCRPage], list[str]]:
    # DOCX files usually contain native text and tables.
    # Embedded images are extracted from word/media/ and OCR is applied to those images.
    warnings: list[str] = []

    try:
        document = Document(docx_path)

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted DOCX file.",
        ) from exc

    chunks: list[str] = []

    # Paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            chunks.append(text)

    # Tables
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\n### Table {table_index}\n")

        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                value = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                cells.append(value)

            if any(cells):
                chunks.append(" | ".join(cells))

    # Embedded images
    image_results = await ocr_docx_embedded_images(docx_path)

    for image_index, image_text, confidence in image_results:
        chunks.append(f"\n### Embedded Image {image_index} OCR\n")

        if image_text.strip():
            chunks.append(image_text.strip())

            if confidence is not None:
                chunks.append(f"\n[OCR confidence: {confidence:.3f}]")
        else:
            chunks.append("[No text detected in embedded image.]")

    text = "\n".join(chunks).strip()

    if not text:
        warnings.append("DOCX file did not contain extractable text.")

    if image_results:
        warnings.append(
            f"DOCX embedded image OCR processed {len(image_results)} image(s)."
        )

    return [
        OCRPage(
            page=1,
            text=text,
            confidence=None,
        )
    ], warnings


async def ocr_docx_embedded_images(
    docx_path: Path,
) -> list[tuple[int, str, float | None]]:
    """Extract images from DOCX and run OCR on supported embedded images."""
    results: list[tuple[int, str, float | None]] = []

    image_suffixes = {
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".tif",
        ".tiff",
        ".webp",
    }

    extracted_dir = settings.upload_dir / f"{docx_path.stem}_docx_images"
    extracted_dir.mkdir(parents=True, exist_ok=True)

    try:
        # DOCX is internally a ZIP file.
        # Embedded images are stored under word/media/.
        with zipfile.ZipFile(docx_path) as archive:
            image_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and Path(name).suffix.lower() in image_suffixes
            ]

            for image_index, image_name in enumerate(image_names, start=1):
                suffix = Path(image_name).suffix.lower()
                image_path = extracted_dir / f"embedded_{image_index}{suffix}"

                image_path.write_bytes(archive.read(image_name))

                try:
                    validate_image(image_path)
                    text, confidence = await get_engine().extract_text(image_path)
                    results.append((image_index, text, confidence))

                except Exception as exc:
                    logger.warning(
                        "Failed to OCR embedded DOCX image %s: %s",
                        image_name,
                        exc,
                    )
                    results.append(
                        (
                            image_index,
                            f"[Failed to OCR embedded image: {exc}]",
                            None,
                        )
                    )

    except zipfile.BadZipFile as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted DOCX file.",
        ) from exc

    finally:
        # removes temporary extracted embedded images
        shutil.rmtree(extracted_dir, ignore_errors=True)

    return results
