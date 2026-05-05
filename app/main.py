import logging
import shutil
import uuid
import zipfile
from pathlib import Path

import fitz
from docx import Document
from fastapi import Depends, FastAPI, File, HTTPException, UploadFile, status
from PIL import Image, UnidentifiedImageError

from app.auth import require_api_key
from app.cache import build_cache_key, get_cached_response, save_cached_response, sha256_file
from app.config import settings
from app.schemas import HealthResponse, OCRPage, OCRResponse
from app.ocr import get_engine


#logger configuration
logger = logging.getLogger("paddleocr_backend")

logging.basicConfig(
    level=getattr(logging, settings.log_level.upper(), logging.INFO),
    format="%(levelname)s:%(name)s:%(message)s",
)


#creates FastAPI app
app = FastAPI(
    title="OpenWebUI PaddleOCR Backend",
    description="OCR and document text extraction backend for Open WebUI.",
    version="0.1.0",
)

#when backend starts
@app.on_event("startup")
async def startup_event() -> None:
    settings.upload_dir.mkdir(parents=True, exist_ok=True) #creates a upload directory for the uploaded images
    logger.info("OCR backend started with engine=%s", settings.ocr_engine)


#GET /health endpoint, returns base config of the ocr backend
@app.get("/health", response_model=HealthResponse)
async def health() -> HealthResponse:
    return HealthResponse(
        status="ok",
        engine=settings.ocr_engine,
        language=settings.ocr_lang,
        max_file_size_mb=settings.max_file_size_mb,
        max_pdf_pages=settings.max_pdf_pages,
    )

#POST /ocr endpoint, 
@app.post("/ocr", response_model=OCRResponse, dependencies=[Depends(require_api_key)]) #checks API key
async def ocr_file(file: UploadFile = File(...)) -> OCRResponse:
    original_name = file.filename or "uploaded-file" #file upload name
    suffix = Path(original_name).suffix.lower() #gets the .extention (".pdf", ".png", ...)

    if suffix not in settings.allowed_extensions: #checks if file type is allowed
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=(
                f"Unsupported file type '{suffix}'. "
                f"Supported: {sorted(settings.allowed_extensions)}"
            ),
        )

    saved_path = await _save_upload(file, suffix) #saves uploaded file in temporary dir
    logger.info("Received OCR request: filename=%s saved=%s", original_name, saved_path)

    try:
        file_hash = sha256_file(saved_path) #saves uploaded files hash
        cache_key = build_cache_key(file_hash) # saves cache key, including file hash and ocr config

        cached_response = get_cached_response(cache_key) 
        if cached_response is not None: #checks if file has been processed before
            logger.info("Returning cached OCR result for filename=%s", original_name)
            cached_response["filename"] = original_name

            warnings = cached_response.get("warnings") or []
            warnings.append("Returned cached OCR result; file was not processed again.")
            cached_response["warnings"] = warnings

            return OCRResponse(**cached_response) 

        if suffix == ".pdf":
            file_type = "pdf"
            pages, warnings = await _ocr_pdf(saved_path)

        elif suffix == ".docx":
            file_type = "docx"
            pages, warnings = await _extract_docx_text(saved_path)

        else:
            file_type = "image"
            _validate_image(saved_path)
            text, confidence = await get_engine().extract_text(saved_path)
            pages = [OCRPage(page=1, text=text, confidence=confidence)]
            warnings = []

        combined_text = "\n\n".join( #joins all pages together in one text
            f"## Page {page.page}\n\n{page.text}" for page in pages
        ).strip()

        if not combined_text: #if the full text is empty throws warning
            warnings.append("OCR completed, but no text was detected.")

        response = OCRResponse( # makes the json ocr response
            filename=original_name,
            file_type=file_type,
            engine=get_engine().name,
            language=settings.ocr_lang,
            page_count=len(pages),
            text=combined_text,
            pages=pages,
            warnings=warnings,
        )

        response_data = (
            response.model_dump()
            if hasattr(response, "model_dump")
            else response.dict()
        )

        try:
            save_cached_response(
                cache_key=cache_key,
                file_hash=file_hash,
                response_data=response_data,
            )
        except Exception:
            logger.warning("Failed to save OCR cache for %s", original_name, exc_info=True)

        return response

    except HTTPException:
        raise

    except Exception as exc:
        logger.exception("OCR processing failed for %s", original_name)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"OCR processing failed: {exc}",
        ) from exc

    finally:
        try:
            saved_path.unlink(missing_ok=True)
        except Exception:
            logger.warning("Could not remove uploaded file: %s", saved_path)


async def _save_upload(file: UploadFile, suffix: str) -> Path:
    settings.upload_dir.mkdir(parents=True, exist_ok=True)

    saved_path = settings.upload_dir / f"{uuid.uuid4().hex}{suffix}"
    max_size_bytes = settings.max_file_size_mb * 1024 * 1024
    total_size = 0

    with saved_path.open("wb") as output:
        while True:
            chunk = await file.read(1024 * 1024)

            if not chunk:
                break

            total_size += len(chunk)

            if total_size > max_size_bytes:
                output.close()
                saved_path.unlink(missing_ok=True)
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=(
                        f"File is too large. Maximum allowed size is "
                        f"{settings.max_file_size_mb} MB."
                    ),
                )

            output.write(chunk)

    return saved_path


def _validate_image(path: Path) -> None:
    try:
        with Image.open(path) as image:
            image.verify()

    except (UnidentifiedImageError, OSError) as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted image file.",
        ) from exc


async def _ocr_pdf(pdf_path: Path) -> tuple[list[OCRPage], list[str]]:
    warnings: list[str] = []
    pages: list[OCRPage] = []
    rendered_dir: Path | None = None

    try:
        doc = fitz.open(pdf_path)

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted PDF file.",
        ) from exc

    native_pages = 0
    ocr_pages = 0

    try:
        total_pages = doc.page_count

        if total_pages == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="PDF file does not contain any pages.",
            )

        if settings.max_pdf_pages <= 0:
            pages_to_process = total_pages
        else:
            pages_to_process = min(total_pages, settings.max_pdf_pages)

            if total_pages > settings.max_pdf_pages:
                warnings.append(
                    f"PDF has {total_pages} pages, but only first "
                    f"{settings.max_pdf_pages} page(s) were processed."
                )

        zoom = settings.pdf_dpi / 72
        matrix = fitz.Matrix(zoom, zoom)

        for page_index in range(pages_to_process):
            page = doc.load_page(page_index)

            native_text = ""
            if settings.pdf_native_text_first:
                native_text = _extract_native_pdf_text(page)

            if _meaningful_char_count(native_text) >= settings.pdf_native_min_chars:
                pages.append(
                    OCRPage(
                        page=page_index + 1,
                        text=native_text,
                        confidence=None,
                    )
                )
                native_pages += 1
                continue

            if rendered_dir is None:
                rendered_dir = settings.upload_dir / f"{pdf_path.stem}_pages"
                rendered_dir.mkdir(parents=True, exist_ok=True)

            image_path = rendered_dir / f"page_{page_index + 1}.png"
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            pix.save(image_path)

            text, confidence = await get_engine().extract_text(image_path)

            pages.append(
                OCRPage(
                    page=page_index + 1,
                    text=text,
                    confidence=confidence,
                )
            )
            ocr_pages += 1

        if settings.pdf_native_text_first:
            warnings.append(
                f"PDF extraction mode: native text used for {native_pages} page(s), "
                f"OCR fallback used for {ocr_pages} page(s)."
            )

        return pages, warnings

    finally:
        doc.close()

        if rendered_dir is not None:
            shutil.rmtree(rendered_dir, ignore_errors=True)


def _extract_native_pdf_text(page) -> str:
    """Extract selectable text from a PDF page using PyMuPDF."""
    try:
        text = page.get_text("text") or ""

    except Exception:
        return ""

    lines: list[str] = []
    previous_blank = False

    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            if lines and not previous_blank:
                lines.append("")
            previous_blank = True
            continue

        lines.append(line)
        previous_blank = False

    return "\n".join(lines).strip()


def _meaningful_char_count(text: str) -> int:
    """Count meaningful alphanumeric characters to decide if native PDF text is usable."""
    return sum(1 for char in text if char.isalnum())


async def _extract_docx_text(docx_path: Path) -> tuple[list[OCRPage], list[str]]:
    warnings: list[str] = []

    try:
        document = Document(docx_path)

    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted DOCX file.",
        ) from exc

    chunks: list[str] = []

    # Paragraphs
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            chunks.append(text)

    # Tables
    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\n### Table {table_index}\n")

        for row in table.rows:
            cells: list[str] = []

            for cell in row.cells:
                value = " ".join(
                    paragraph.text.strip()
                    for paragraph in cell.paragraphs
                    if paragraph.text.strip()
                )
                cells.append(value)

            if any(cells):
                chunks.append(" | ".join(cells))

    # Embedded images
    image_results = await _ocr_docx_embedded_images(docx_path)

    for image_index, image_text, confidence in image_results:
        chunks.append(f"\n### Embedded Image {image_index} OCR\n")

        if image_text.strip():
            chunks.append(image_text.strip())

            if confidence is not None:
                chunks.append(f"\n[OCR confidence: {confidence:.3f}]")
        else:
            chunks.append("[No text detected in embedded image.]")

    text = "\n".join(chunks).strip()

    if not text:
        warnings.append("DOCX file did not contain extractable text.")

    if image_results:
        warnings.append(
            f"DOCX embedded image OCR processed {len(image_results)} image(s)."
        )

    return [
        OCRPage(
            page=1,
            text=text,
            confidence=None,
        )
    ], warnings


async def _ocr_docx_embedded_images(
    docx_path: Path,
) -> list[tuple[int, str, float | None]]:
    """Extract images from DOCX and run OCR on supported embedded images."""
    results: list[tuple[int, str, float | None]] = []

    image_suffixes = {
        ".png",
        ".jpg",
        ".jpeg",
        ".bmp",
        ".tif",
        ".tiff",
        ".webp",
    }

    extracted_dir = settings.upload_dir / f"{docx_path.stem}_docx_images"
    extracted_dir.mkdir(parents=True, exist_ok=True)

    try:
        with zipfile.ZipFile(docx_path) as archive:
            image_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/")
                and Path(name).suffix.lower() in image_suffixes
            ]

            for image_index, image_name in enumerate(image_names, start=1):
                suffix = Path(image_name).suffix.lower()
                image_path = extracted_dir / f"embedded_{image_index}{suffix}"

                image_path.write_bytes(archive.read(image_name))

                try:
                    _validate_image(image_path)
                    text, confidence = await get_engine().extract_text(image_path)
                    results.append((image_index, text, confidence))

                except Exception as exc:
                    logger.warning(
                        "Failed to OCR embedded DOCX image %s: %s",
                        image_name,
                        exc,
                    )
                    results.append(
                        (
                            image_index,
                            f"[Failed to OCR embedded image: {exc}]",
                            None,
                        )
                    )

    except zipfile.BadZipFile as exc:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid or corrupted DOCX file.",
        ) from exc

    finally:
        shutil.rmtree(extracted_dir, ignore_errors=True)

    return results
