from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches


BASE_DIR = Path("samples")
IMG_DIR = BASE_DIR / "images"
PDF_DIR = BASE_DIR / "pdf"
DOCX_DIR = BASE_DIR / "docx"
TMP_DIR = BASE_DIR / "_generated_tmp"

for directory in [IMG_DIR, PDF_DIR, DOCX_DIR, TMP_DIR]:
    directory.mkdir(parents=True, exist_ok=True)

FONT_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD_PATH = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
FONT_NAME = "DejaVuSans"

pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))


EN_LINES = [
    "English OCR Test Sample",
    "",
    "This sample is used to validate image OCR, PDF extraction, and DOCX extraction.",
    "It contains headings, paragraphs, numbers, punctuation, and a small table.",
    "",
    "Use case:",
    "- Extract text from uploaded documents.",
    "- Preserve page separation.",
    "- Support image-based OCR fallback.",
    "",
    "Invoice: INV-2026-0042",
    "Date: 2026-04-29",
    "Total: EUR 125.50",
    "",
    "Table:",
    "Item | Quantity | Status",
    "Book | 2 | Scanned",
    "Notebook | 5 | Ready",
    "",
    "Expected result: readable English text with correct structure.",
]

LV_LINES = [
    "Latviešu OCR testa paraugs",
    "",
    "Šis paraugs tiek izmantots attēlu OCR, PDF un DOCX izvilkšanas pārbaudei.",
    "Tas satur virsrakstus, rindkopas, ciparus, pieturzīmes un nelielu tabulu.",
    "",
    "Lietošanas piemērs:",
    "- Izvilkt tekstu no augšupielādētiem dokumentiem.",
    "- Saglabāt lapu sadalījumu.",
    "- Atbalstīt OCR skenētām vai attēla tipa lapām.",
    "",
    "Rēķins: RIK-2026-0042",
    "Datums: 2026-04-29",
    "Summa: EUR 125,50",
    "",
    "Tabula:",
    "Prece | Daudzums | Statuss",
    "Grāmata | 2 | Skenēts",
    "Piezīmju klade | 5 | Gatavs",
    "",
    "Sagaidāmais rezultāts: salasāms latviešu teksts ar garumzīmēm un mīkstinājumiem.",
]


def create_text_image(lines: list[str], output_path: Path) -> None:
    width, height = 1700, 2200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = ImageFont.truetype(FONT_BOLD_PATH, 56)
    body_font = ImageFont.truetype(FONT_PATH, 38)

    x = 120
    y = 110

    for index, line in enumerate(lines):
        if not line:
            y += 34
            continue

        font = title_font if index == 0 else body_font
        draw.text((x, y), line, fill="black", font=font)
        y += 74 if index == 0 else 54

    image.save(output_path)


def create_image_samples() -> None:
    en_png = IMG_DIR / "english_image_test.png"
    lv_png = IMG_DIR / "latvian_image_test.png"

    create_text_image(EN_LINES, en_png)
    create_text_image(LV_LINES, lv_png)

    Image.open(en_png).save(IMG_DIR / "english_image_test.jpg", quality=95)
    Image.open(lv_png).save(IMG_DIR / "latvian_image_test.jpg", quality=95)


def draw_digital_pdf_page(pdf: canvas.Canvas, lines: list[str]) -> None:
    width, height = A4
    x = 24 * mm
    y = height - 25 * mm

    pdf.setFont(FONT_NAME, 17)
    pdf.drawString(x, y, lines[0])
    y -= 14 * mm

    pdf.setFont(FONT_NAME, 10.5)

    for line in lines[1:]:
        if not line:
            y -= 5 * mm
            continue

        pdf.drawString(x, y, line)
        y -= 6.5 * mm

        if y < 25 * mm:
            pdf.showPage()
            pdf.setFont(FONT_NAME, 10.5)
            y = height - 25 * mm


def create_mixed_pdf(lines: list[str], output_path: Path) -> None:
    scanned_image = TMP_DIR / f"{output_path.stem}_scanned_page.png"
    create_text_image(lines, scanned_image)

    pdf = canvas.Canvas(str(output_path), pagesize=A4)

    # Page 1: real selectable PDF text
    draw_digital_pdf_page(pdf, lines)
    pdf.showPage()

    # Page 2: image-only page to force OCR fallback
    pdf.drawImage(str(scanned_image), 0, 0, width=A4[0], height=A4[1])
    pdf.showPage()

    pdf.save()


def create_pdf_samples() -> None:
    create_mixed_pdf(EN_LINES, PDF_DIR / "english_mixed_test.pdf")
    create_mixed_pdf(LV_LINES, PDF_DIR / "latvian_mixed_test.pdf")


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"

    for idx, value in enumerate(rows[0]):
        table.rows[0].cells[idx].text = value

    for row_values in rows[1:]:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            row.cells[idx].text = value


def create_docx_sample(lines: list[str], image_path: Path, output_path: Path, language: str) -> None:
    document = Document()

    document.add_heading(f"{language} DOCX Embedded Image OCR Test", level=1)
    document.add_paragraph(
        "This document validates native DOCX text extraction, table extraction, "
        "and OCR for an image embedded inside the DOCX file."
    )

    document.add_heading("Native Text Section", level=2)

    if language == "English":
        document.add_paragraph(
            "English native text: OpenWebUI PaddleOCR Backend should extract this paragraph without OCR."
        )
    else:
        document.add_paragraph(
            "Latviešu native teksts: sistēmai jāizvelk šī rindkopa bez OCR, saglabājot ā, ē, ī, ū, č, ģ, ķ, ļ, ņ, š, ž."
        )

    document.add_heading("Native Table Section", level=2)

    if language == "English":
        add_docx_table(
            document,
            [
                ["ID", "Language", "Expected extraction", "Status"],
                ["T-001", "English", "Native paragraph text", "Should pass"],
                ["T-002", "English image", "OCR text from embedded image", "Should pass"],
            ],
        )
    else:
        add_docx_table(
            document,
            [
                ["ID", "Valoda", "Sagaidāmā izvilkšana", "Statuss"],
                ["T-001", "Latviešu", "Native rindkopas teksts", "Jāizdodas"],
                ["T-002", "Latviešu bilde", "OCR teksts no iegultas bildes", "Jāizdodas"],
            ],
        )

    document.add_heading("Embedded Image OCR Section", level=2)
    document.add_paragraph(
        "The image below contains text and should be extracted from word/media/ and processed with OCR."
    )
    document.add_picture(str(image_path), width=Inches(5.8))

    document.add_paragraph("End of sample document.")
    document.save(output_path)


def create_docx_samples() -> None:
    en_image = TMP_DIR / "english_docx_embedded_image.png"
    lv_image = TMP_DIR / "latvian_docx_embedded_image.png"

    create_text_image(EN_LINES, en_image)
    create_text_image(LV_LINES, lv_image)

    create_docx_sample(
        EN_LINES,
        en_image,
        DOCX_DIR / "english_docx_embedded_image_test.docx",
        "English",
    )

    create_docx_sample(
        LV_LINES,
        lv_image,
        DOCX_DIR / "latvian_docx_embedded_image_test.docx",
        "Latvian",
    )


def main() -> None:
    create_image_samples()
    create_pdf_samples()
    create_docx_samples()

    for file in sorted(BASE_DIR.glob("*/*")):
        if file.is_file():
            print(file)


if __name__ == "__main__":
    main()
